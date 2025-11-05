#!/usr/bin/env python3
"""
Calendar tool for organizing lab protocols.
Parses PDF or DOC/DOCX files for Day X entries and assigns calendar dates.

Standalone script - no virtual environment required.
Install dependencies with: pip3 install -r requirements.txt
"""

import argparse
import re
from datetime import datetime, timedelta
from pathlib import Path
import sys
import platform

# Import platform-agnostic calendar functions
try:
    from calendar_platform import (
        add_to_calendar as add_to_apple_calendar,
        remove_from_calendar as remove_from_apple_calendar,
        find_matching_experiment_ids,
        extract_day0_from_events
    )
except ImportError:
    # Fallback for backward compatibility (will use platform-specific functions if available)
    def add_to_apple_calendar(*args, **kwargs):
        print("Error: Calendar backend not available. Install platform-specific dependencies.")
        return False
    def remove_from_apple_calendar(*args, **kwargs):
        print("Error: Calendar backend not available. Install platform-specific dependencies.")
        return False
    def find_matching_experiment_ids(*args, **kwargs):
        return {}
    def extract_day0_from_events(*args, **kwargs):
        return None

# Export functions for use by GUI
__all__ = [
    'extract_text_from_pdf', 'extract_text_from_docx', 'extract_text_from_doc',
    'parse_day_entries', 'parse_date', 'assign_dates', 'extract_title_and_id',
    'add_to_apple_calendar', 'remove_from_apple_calendar',
    'find_matching_experiment_ids', 'extract_day0_from_events'
]


def extract_text_from_pdf(pdf_path):
    """Extract text from PDF file."""
    try:
        import PyPDF2
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    except ImportError:
        try:
            import pdfplumber
            text = ""
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            return text
        except ImportError:
            print("Error: Need either PyPDF2 or pdfplumber to read PDF files.")
            print("Install with: pip install PyPDF2 or pip install pdfplumber")
            sys.exit(1)


def extract_text_from_docx(docx_path):
    """Extract text from DOCX file."""
    try:
        from docx import Document
        doc = Document(docx_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    except ImportError:
        print("Error: Need python-docx to read DOCX files.")
        print("Install with: pip install python-docx")
        sys.exit(1)


def extract_text_from_doc(doc_path):
    """Extract text from DOC file (older format)."""
    try:
        import docx2python
        result = docx2python.docx2python(doc_path)
        return result.text
    except ImportError:
        # Fallback: try python-docx (might work for some DOC files)
        try:
            from docx import Document
            doc = Document(doc_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text
        except Exception as e:
            print(f"Error reading DOC file: {e}")
            print("Try installing: pip install docx2python or convert DOC to DOCX")
            sys.exit(1)


def parse_day_entries(text):
    """
    Parse Day X entries from text.
    Extracts text after "Day X:" up to the first period.
    Returns a list of tuples: (day_number, task_description)
    """
    # Pattern to match "Day X:" or "Day X-Y:" followed by description until first period
    # Handles ranges like "Day 5-6:" by taking the first number
    pattern = r'Day\s*(\d+)(?:-\d+)?:\s*([^.]*?\.)'
    
    matches = re.finditer(pattern, text, re.IGNORECASE | re.DOTALL)
    
    entries = []
    for match in matches:
        day_num = int(match.group(1))
        task = match.group(2).strip()
        # Remove the trailing period and clean up whitespace
        task = task.rstrip('.').strip()
        task = re.sub(r'\s+', ' ', task)
        if task:  # Only add non-empty tasks
            entries.append((day_num, task))
    
    # Sort by day number
    entries.sort(key=lambda x: x[0])
    
    return entries


def extract_title_and_id(text):
    """
    Extract title from text and generate ID from first 3 letters.
    Returns tuple: (title, id)
    """
    # Get first non-empty line as title
    lines = text.strip().split('\n')
    title = None
    for line in lines:
        line = line.strip()
        if line and not line.startswith('Day'):
            title = line
            break
    
    if title:
        # Get first 3 letters (alphanumeric), uppercase
        id_text = ''.join(c for c in title if c.isalnum())[:3].upper()
        return title, id_text
    
    return None, None


def parse_date(date_str):
    """
    Parse date string in various formats.
    Supports: MM/DD/YY, MM/DD/YYYY, YYYY-MM-DD, etc.
    """
    formats = [
        '%m/%d/%y',    # 10/11/25
        '%m/%d/%Y',    # 10/11/2025
        '%Y-%m-%d',    # 2025-10-11
        '%m-%d-%Y',    # 10-11-2025
        '%m-%d-%y',    # 10-11-25
    ]
    
    for fmt in formats:
        try:
            return datetime.strptime(date_str, fmt)
        except ValueError:
            continue
    
    raise ValueError(f"Could not parse date: {date_str}. Supported formats: MM/DD/YY, MM/DD/YYYY, YYYY-MM-DD")


# Cache for spaCy model
_nlp_model = None

def get_spacy_model():
    """Get or load spaCy model (cached)."""
    global _nlp_model
    if _nlp_model is None:
        try:
            import spacy
            try:
                _nlp_model = spacy.load("en_core_web_sm")
            except OSError:
                # Model not found
                return None
        except ImportError:
            return None
    return _nlp_model


def simplify_task_nlp(task_text):
    """
    Simplify task description to 2-3 words using NLP.
    Tries spaCy first, falls back to keyword extraction.
    """
    nlp = get_spacy_model()
    
    if nlp is not None:
        doc = nlp(task_text)
        
        # Extract important verbs and nouns
        keywords = []
        for token in doc:
            # Get verbs (action words) and important nouns
            if token.pos_ in ['VERB', 'NOUN'] and not token.is_stop:
                # Skip very common words
                if token.text.lower() not in ['the', 'a', 'an', 'is', 'are', 'was', 'were']:
                    keywords.append(token.lemma_.lower())
        
        # Also extract named entities (lab terms, etc.)
        for ent in doc.ents:
            if ent.label_ in ['PERSON', 'ORG', 'PRODUCT', 'EVENT']:
                keywords.append(ent.text.lower())
        
        # If we found keywords, take first 2-3
        if keywords:
            simplified = ' '.join(keywords[:3])
            return simplified
    
    # Fallback to heuristic method
    return simplify_task_heuristic(task_text)


def simplify_task_heuristic(task_text):
    """
    Simplify task using simple heuristics when NLP library unavailable.
    Extracts key action verbs and important nouns.
    """
    # Common lab protocol action words
    lab_actions = [
        'plate', 'seed', 'treat', 'harvest', 'collect', 'analyze', 'measure',
        'change', 'replace', 'add', 'remove', 'wash', 'stain', 'fix', 'count',
        'incubate', 'culture', 'split', 'passage', 'freeze', 'thaw', 'lyse',
        'extract', 'isolate', 'purify', 'centrifuge', 'spin', 'pipette', 'dilute',
        'prepare', 'set', 'start', 'begin', 'end', 'finish', 'check', 'verify'
    ]
    
    words = task_text.lower().split()
    
    # Remove common stop words and articles
    stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 
                  'of', 'with', 'by', 'from', 'as', 'is', 'are', 'was', 'were', 'be',
                  'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will',
                  'would', 'should', 'could', 'may', 'might', 'must', 'can'}
    
    keywords = []
    
    # Look for action verbs first
    for word in words:
        # Clean word (remove punctuation)
        clean_word = re.sub(r'[^\w\s]', '', word)
        if clean_word in lab_actions and clean_word not in keywords:
            keywords.append(clean_word)
            if len(keywords) >= 2:
                break
    
    # Add important nouns if we don't have enough keywords
    for word in words:
        clean_word = re.sub(r'[^\w\s]', '', word)
        if (clean_word not in stop_words and 
            clean_word not in keywords and 
            len(clean_word) > 3):  # Skip short words
            keywords.append(clean_word)
            if len(keywords) >= 3:
                break
    
    if keywords:
        return ' '.join(keywords[:3])
    
    # Final fallback: take first 2-3 non-stop words
    fallback = [w for w in words if w not in stop_words and len(w) > 2][:3]
    return ' '.join(fallback) if fallback else task_text[:30]  # Last resort: truncate


def simplify_tasks(entries, use_nlp=True):
    """
    Simplify all task descriptions to 2-3 words.
    Returns entries with simplified tasks: (day_number, task_description, simplified_task)
    """
    simplified_entries = []
    for day_num, task in entries:
        if use_nlp:
            simplified = simplify_task_nlp(task)
        else:
            simplified = simplify_task_heuristic(task)
        simplified_entries.append((day_num, task, simplified))
    return simplified_entries


def assign_dates(entries, day0_date, exp_id=None):
    """
    Assign calendar dates to day entries.
    Returns a list of tuples: (day_number, task_description, calendar_date, exp_id)
    """
    dated_entries = []
    for day_num, task in entries:
        calendar_date = day0_date + timedelta(days=day_num)
        dated_entries.append((day_num, task, calendar_date, exp_id))
    
    return dated_entries


def print_calendar(dated_entries):
    """Print the calendar schedule in a readable format."""
    print("\n" + "="*80)
    print("LAB PROTOCOL CALENDAR")
    print("="*80)
    print()
    
    for day_num, task, calendar_date, exp_id in dated_entries:
        date_str = calendar_date.strftime("%A, %B %d, %Y")
        
        if exp_id:
            print(f"Day {day_num:3d} ({date_str}):")
            print(f"  ID: {exp_id}, Day {day_num}: {task}")
        else:
            print(f"Day {day_num:3d} ({date_str}):")
            print(f"  {task}")
        print()


# Calendar functions are now imported from calendar_platform.py at the top of this file
# This provides platform-agnostic support (macOS EventKit or universal iCalendar)


def main():
    parser = argparse.ArgumentParser(
        description="Parse lab protocol files and assign calendar dates to Day entries",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  # Parse protocol and automatically add to Apple Calendar (when --id is provided)
  python calendar_parser.py protocol.pdf --id ENC
  
  # Parse with custom Day 0 date
  python calendar_parser.py protocol.docx --day0 10/11/25 --id EXP001
  
  # Parse only (don't add to calendar)
  python calendar_parser.py protocol.pdf --id ENC --no-calendar
  
  # Remove events from Apple Calendar by experiment ID
  python calendar_parser.py --remove-from-calendar --id ENC
        """
    )
    
    parser.add_argument('file', 
                       nargs='?',
                       help='Path to PDF or DOC/DOCX file containing protocol (not needed for --remove-from-calendar)')
    parser.add_argument('--day0', 
                       type=str, 
                       default=None,
                       help='Start date for Day 0 (e.g., 10/11/25, 10/11/2025, or 2025-10-11). Defaults to today.')
    parser.add_argument('--id',
                       type=str,
                       default=None,
                       help='Experiment ID to label all calendar entries (e.g., EXP001, ENC). Automatically adds to calendar when provided.')
    parser.add_argument('--no-calendar',
                       action='store_true',
                       help='Skip adding events to Apple Calendar (default: automatically adds when --id is provided).')
    parser.add_argument('--remove-from-calendar',
                       action='store_true',
                       help='Remove events from Apple Calendar by experiment ID.')
    parser.add_argument('--calendar-name',
                       type=str,
                       default='Lab Protocols',
                       help='Name of the calendar to use (default: Lab Protocols).')
    
    args = parser.parse_args()
    
    # If removing from calendar, we only need the ID
    if args.remove_from_calendar:
        if not args.id:
            print("Error: --id is required when removing events from calendar.")
            sys.exit(1)
        success = remove_from_apple_calendar(args.id, args.calendar_name)
        sys.exit(0 if success else 1)
    
    # For other operations, file is required
    if not args.file:
        print("Error: File path is required (unless using --remove-from-calendar).")
        parser.print_help()
        sys.exit(1)
    
    # Auto-add to calendar if ID is provided (unless --no-calendar is set)
    should_add_to_calendar = args.id and not args.no_calendar
    
    # Check if file exists
    file_path = Path(args.file)
    if not file_path.exists():
        print(f"Error: File not found: {args.file}")
        sys.exit(1)
    
    # Determine file type and extract text
    file_ext = file_path.suffix.lower()
    print(f"Reading file: {args.file}...")
    
    if file_ext == '.pdf':
        text = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        text = extract_text_from_docx(file_path)
    elif file_ext == '.doc':
        text = extract_text_from_doc(file_path)
    else:
        print(f"Error: Unsupported file type: {file_ext}")
        print("Supported formats: .pdf, .doc, .docx")
        sys.exit(1)
    
    # Extract title (for display purposes)
    title, _ = extract_title_and_id(text)
    if title:
        print(f"Title: {title}")
    
    # Use provided ID (may be None if not adding to calendar)
    exp_id = args.id
    
    # Parse day entries
    print("Parsing Day entries...")
    entries = parse_day_entries(text)
    
    if not entries:
        print("Warning: No Day entries found in the document.")
        print("Looking for patterns like 'Day 0:', 'Day 1:', etc.")
        sys.exit(1)
    
    print(f"Found {len(entries)} day entries.")
    
    # Parse Day 0 date
    if args.day0:
        try:
            day0_date = parse_date(args.day0)
            # Set time to midnight
            day0_date = day0_date.replace(hour=0, minute=0, second=0, microsecond=0)
        except ValueError as e:
            print(f"Error: {e}")
            sys.exit(1)
    else:
        day0_date = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
    
    print(f"Day 0 set to: {day0_date.strftime('%A, %B %d, %Y')}")
    
    if exp_id:
        print(f"Experiment ID: {exp_id}")
    
    # Assign dates
    dated_entries = assign_dates(entries, day0_date, exp_id)
    
    # Print calendar
    print_calendar(dated_entries)
    
    # Optionally save to file
    if exp_id:
        output_file = f"{file_path.stem}_{exp_id}_calendar.txt"
    else:
        output_file = file_path.stem + '_calendar.txt'
    
    with open(output_file, 'w') as f:
        f.write(f"LAB PROTOCOL CALENDAR\n")
        if title:
            f.write(f"Title: {title}\n")
        if exp_id:
            f.write(f"Experiment ID: {exp_id}\n")
        f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        f.write(f"Day 0: {day0_date.strftime('%A, %B %d, %Y')}\n")
        f.write("="*80 + "\n\n")
        for day_num, task, calendar_date, exp_id_entry in dated_entries:
            date_str = calendar_date.strftime("%A, %B %d, %Y")
            
            if exp_id_entry:
                f.write(f"Day {day_num:3d} ({date_str}):\n")
                f.write(f"  ID: {exp_id_entry}, Day {day_num}: {task}\n\n")
            else:
                f.write(f"Day {day_num:3d} ({date_str}):\n")
                f.write(f"  {task}\n\n")
    
    print(f"\nCalendar saved to: {output_file}")
    
    # Automatically add to Apple Calendar if ID is provided (unless --no-calendar)
    if should_add_to_calendar:
        print("\nAdding events to Apple Calendar...")
        add_to_apple_calendar(dated_entries, exp_id, args.calendar_name)


if __name__ == '__main__':
    main()

